# -*- coding: utf-8 -*-
"""M8 PLUV - Gerador do Memorial Pluvial NBR 10844 / DTU 60.11.
Gera relatorios em HTML, PDF e DOCX com dados dinamicos do projeto (IronPython e CPython).
"""
import codecs
import json
import os
from datetime import datetime

try:
    from xhtml2pdf import pisa
except ImportError:
    pisa = None

try:
    from docx import Document
except ImportError:
    Document = None

_this_dir = os.path.dirname(os.path.abspath(__file__))
_auto_root = os.path.dirname(_this_dir) if os.path.basename(_this_dir) == "tools" else _this_dir
RAIZ = globals().get("RAIZ", os.environ.get("HYDRO_PROJECT_ROOT", _auto_root))
D = os.path.join(RAIZ, "data")
SAIDA = os.path.join(RAIZ, "memoriais")


def ler_json(caminho):
    if not os.path.isfile(caminho):
        return {}
    f = codecs.open(caminho, "r", encoding="utf-8")
    dados = json.loads(f.read())
    f.close()
    return dados


R_PLUV = ler_json(os.path.join(D, "dimensionamento_pluv.json"))
CFGP = ler_json(os.path.join(D, "config_projeto.json"))
RESP = CFGP.get("responsavel_tecnico", {"nome": "Thayná Barreiro", "titulo": "Engenheira / Desenhista"})

proj = CFGP.get("projeto", {})
proprietario = proj.get("proprietario", proj.get("nome", "Cliente do Projeto"))
nome_projeto = proj.get("nome", "Residência Unifamiliar")
localizacao = proj.get("cidade", "Porto Alegre / SFS")

d_pluv = R_PLUV.get("pluvial", {})

CSS = """
@page { size: A4; margin: 0; }
* { box-sizing: border-box; }
body { font-family: 'Segoe UI', Arial, sans-serif; font-size: 10pt; color: #222222; line-height: 1.5; margin: 0; padding: 0; background: #ffffff; }

.page { width: 210mm; min-height: 297mm; padding: 20mm 18mm; margin: 0 auto; page-break-after: always; position: relative; background: #ffffff; }
.page-cover { background: #2b2b2b; color: #ffffff; padding: 0; display: flex; flex-direction: column; justify-content: space-between; height: 297mm; }

.cover-container { height: 100%; padding: 40px; position: relative; display: flex; flex-direction: column; justify-content: space-between; }
.cover-top-line { width: 2px; height: 120px; background: #ffffff; position: absolute; top: 0; right: 80px; }
.cover-frame { border: 2px solid #ffffff; border-right: none; padding: 60px 40px; margin-top: 100px; width: 85%; position: relative; }
.cover-title { font-size: 52pt; font-weight: 700; font-family: Arial, sans-serif; margin: 0; letter-spacing: -1px; }
.cover-subtitle { font-size: 22pt; margin-top: 5px; color: #eeeeee; display: flex; align-items: center; gap: 10px; }
.cover-subline { width: 260px; height: 3px; background: #ffffff; margin-top: 15px; }
.cover-owner { font-size: 18pt; font-weight: 600; margin-top: 80px; color: #ffffff; }
.cover-bottom { border-top: 2px solid #ffffff; padding-top: 15px; display: flex; justify-content: space-between; align-items: center; }

h1.sec-title { font-size: 14pt; color: #000000; text-transform: uppercase; margin-top: 25px; margin-bottom: 12px; font-weight: 700; border-bottom: 2px solid #333333; padding-bottom: 3px; }
h2.sec-subtitle { font-size: 11pt; color: #333333; margin-top: 16px; margin-bottom: 8px; font-weight: 700; }

table.tb-info { width: 100%; border-collapse: collapse; margin: 12px 0 20px; }
table.tb-info th { background: #333333; color: #ffffff; padding: 8px; text-align: center; font-size: 11pt; font-weight: 700; }
table.tb-info td { border: 1px solid #cccccc; padding: 7px 12px; font-size: 10pt; }
table.tb-info td:first-child { font-weight: 600; width: 40%; background: #f9f9f9; }

table.tb-data { width: 100%; border-collapse: collapse; margin: 10px 0 16px; font-size: 9.5pt; }
table.tb-data th { background: #333333; color: #ffffff; padding: 6px 8px; text-align: left; font-weight: 600; }
table.tb-data td { border: 1px solid #dddddd; padding: 5px 8px; }
table.tb-data tr:nth-child(even) td { background: #fcfcfc; }
"""

h = []
a = h.append

a('<!doctype html><html lang="pt-BR"><head><meta charset="utf-8"><title>Memorial Pluvial — ' + nome_projeto + '</title><style>' + CSS + '</style></head><body>')

# COVER
a('<div class="page page-cover"><div class="cover-container"><div class="cover-top-line"></div>')
a('<div class="cover-frame"><h1 class="cover-title">Memorial</h1><div class="cover-subtitle">Pluvial</div><div class="cover-subline"></div>')
a('<div class="cover-owner">' + proprietario + '</div></div>')
a('<div class="cover-bottom"><span style="color:#aaaaaa;font-size:9pt;">Instalações Prediais de Águas Pluviais</span><div style="font-size:16pt;color:#f1c40f;font-weight:bold;">⌂</div></div>')
a('</div></div>')

# PAGE 2: INFORMAÇÕES
a('<div class="page"><h1 class="sec-title">1 INFORMAÇÕES DO PROJETO</h1>')
a('<table class="tb-info"><tr><th colspan="2">INFORMAÇÕES GERAIS</th></tr>')
a('<tr><td>Empreendimento</td><td>' + nome_projeto + '</td></tr>')
a('<tr><td>Proprietário / Cliente</td><td>' + proprietario + '</td></tr>')
a('<tr><td>Localização / Cidade</td><td>' + localizacao + '</td></tr>')
a('<tr><td>Número de pavimentos</td><td>2 Pavimentos</td></tr>')
a('<tr><td>Tipo de Edificação</td><td>Residencial Unifamiliar</td></tr>')
a('</table>')
a('<p>Profissional Desenhista: <b>' + RESP["nome"] + '</b> (' + RESP["titulo"] + ')</p>')

a('<h1 class="sec-title">2 NORMAS UTILIZADAS</h1>')
a('<ul><li><b>NBR 10844:1989</b> – Instalações Prediais de Águas Pluviais.</li><li><b>DTU 60.11 / NF EN 12056</b> – Évacuation des eaux pluviales.</li></ul>')

a('<h1 class="sec-title">3 DADOS METEOROLÓGICOS E DIMENSIONAMENTO</h1>')
a('<table class="tb-data"><tr><th>Parâmetro</th><th>Valor Adotado</th></tr>')
a('<tr><td>Cidade / Base Pluviométrica</td><td>' + d_pluv.get("cidade", localizacao) + '</td></tr>')
a('<tr><td>Intensidade de Chuva (i)</td><td><b>' + str(d_pluv.get("intensidade_mm_h", 156.0)) + ' mm/h</b></td></tr>')
a('<tr><td>Área Total de Contribuição</td><td>150.00 m²</td></tr>')
a('<tr><td>Vazão Total de Projeto (Q)</td><td><b>' + str(d_pluv.get("vazao_total_ls", 6.5)) + ' L/s (' + str(round(d_pluv.get("vazao_total_ls", 6.5) * 60, 2)) + ' L/min)</b></td></tr>')
a('<tr><td>Condutores Verticais (Prumadas)</td><td>2 x DN 100 mm</td></tr>')
a('<tr><td>Coletor Horizontal Enterrado</td><td>DN 150 mm (0.5% declividade)</td></tr>')
a('</table></div></body></html>')

if not os.path.exists(SAIDA):
    os.makedirs(SAIDA)

nome_base = "Memorial_Pluvial_" + proprietario.replace(" ", "_").replace("&", "e")

# 1. HTML
html_path = os.path.join(SAIDA, nome_base + ".html")
f_html = codecs.open(html_path, "w", encoding="utf-8")
f_html.write("\n".join(h))
f_html.close()
print("HTML gerado:", html_path)

# 2. PDF
if pisa:
    try:
        pdf_path = os.path.join(SAIDA, nome_base + ".pdf")
        f_in = codecs.open(html_path, "r", encoding="utf-8")
        html_text = f_in.read()
        f_in.close()
        f_out = open(pdf_path, "wb")
        pisa.CreatePDF(html_text, dest=f_out)
        f_out.close()
        print("PDF gerado:", pdf_path)
    except Exception as ex_pdf:
        print("Aviso ao gerar PDF:", ex_pdf)

# 3. DOCX
if Document:
    docx_path = os.path.join(SAIDA, nome_base + ".docx")
    doc = Document()
    doc.add_heading("Memorial Pluvial — " + nome_projeto, level=0)
    doc.add_paragraph("Proprietário: " + proprietario)
    doc.add_paragraph("Localização: " + localizacao)
    doc.save(docx_path)
    print("DOCX gerado:", docx_path)
