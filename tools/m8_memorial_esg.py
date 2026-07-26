# -*- coding: utf-8 -*-
"""M8 ESG - Gerador do Memorial Sanitario (Esgoto, Fossa, Filtro, Sumidouro) NBR 8160/7229/13969.
Gera relatorios em HTML, PDF e DOCX com dados dinamicos do projeto (IronPython e CPython).
"""
import codecs
import json
import os
from datetime import datetime

try:
    from xhtml2pdf import pisa
except ImportError:
    pisa = None

try:
    from docx import Document
except ImportError:
    Document = None

_this_dir = os.path.dirname(os.path.abspath(__file__))
_auto_root = os.path.dirname(_this_dir) if os.path.basename(_this_dir) == "tools" else _this_dir
RAIZ = globals().get("RAIZ", os.environ.get("HYDRO_PROJECT_ROOT", _auto_root))
D = os.path.join(RAIZ, "data")
SAIDA = os.path.join(RAIZ, "memoriais")


def ler_json(caminho):
    if not os.path.isfile(caminho):
        return {}
    f = codecs.open(caminho, "r", encoding="utf-8")
    dados = json.loads(f.read())
    f.close()
    return dados


R_ESG = ler_json(os.path.join(D, "dimensionamento_esg.json"))
R_TRAT = ler_json(os.path.join(D, "dimensionamento_trat.json"))
CFGP = ler_json(os.path.join(D, "config_projeto.json"))
RESP = CFGP.get("responsavel_tecnico", {"nome": "Thayná Barreiro", "titulo": "Engenheira / Desenhista"})

proj = CFGP.get("projeto", {})
proprietario = proj.get("proprietario", proj.get("nome", "Cliente do Projeto"))
nome_projeto = proj.get("nome", "Residência Unifamiliar")
localizacao = proj.get("cidade", "Porto Alegre / SFS")

fossa = R_TRAT.get("fossa_septica", {})
filtro = R_TRAT.get("filtro_anaerobio", {})
sumidouro = R_TRAT.get("sumidouro", {})

CSS = """
@page { size: A4; margin: 0; }
* { box-sizing: border-box; }
body { font-family: 'Segoe UI', Arial, sans-serif; font-size: 10pt; color: #222222; line-height: 1.5; margin: 0; padding: 0; background: #ffffff; }

.page { width: 210mm; min-height: 297mm; padding: 20mm 18mm; margin: 0 auto; page-break-after: always; position: relative; background: #ffffff; }
.page-cover { background: #2b2b2b; color: #ffffff; padding: 0; display: flex; flex-direction: column; justify-content: space-between; height: 297mm; }

.cover-container { height: 100%; padding: 40px; position: relative; display: flex; flex-direction: column; justify-content: space-between; }
.cover-top-line { width: 2px; height: 120px; background: #ffffff; position: absolute; top: 0; right: 80px; }
.cover-frame { border: 2px solid #ffffff; border-right: none; padding: 60px 40px; margin-top: 100px; width: 85%; position: relative; }
.cover-title { font-size: 52pt; font-weight: 700; font-family: Arial, sans-serif; margin: 0; letter-spacing: -1px; }
.cover-subtitle { font-size: 22pt; margin-top: 5px; color: #eeeeee; display: flex; align-items: center; gap: 10px; }
.cover-subline { width: 260px; height: 3px; background: #ffffff; margin-top: 15px; }
.cover-owner { font-size: 18pt; font-weight: 600; margin-top: 80px; color: #ffffff; }
.cover-bottom { border-top: 2px solid #ffffff; padding-top: 15px; display: flex; justify-content: space-between; align-items: center; }

h1.sec-title { font-size: 14pt; color: #000000; text-transform: uppercase; margin-top: 25px; margin-bottom: 12px; font-weight: 700; border-bottom: 2px solid #333333; padding-bottom: 3px; }
h2.sec-subtitle { font-size: 11pt; color: #333333; margin-top: 16px; margin-bottom: 8px; font-weight: 700; }

table.tb-info { width: 100%; border-collapse: collapse; margin: 12px 0 20px; }
table.tb-info th { background: #333333; color: #ffffff; padding: 8px; text-align: center; font-size: 11pt; font-weight: 700; }
table.tb-info td { border: 1px solid #cccccc; padding: 7px 12px; font-size: 10pt; }
table.tb-info td:first-child { font-weight: 600; width: 40%; background: #f9f9f9; }

table.tb-data { width: 100%; border-collapse: collapse; margin: 10px 0 16px; font-size: 9.5pt; }
table.tb-data th { background: #333333; color: #ffffff; padding: 6px 8px; text-align: left; font-weight: 600; }
table.tb-data td { border: 1px solid #dddddd; padding: 5px 8px; }
table.tb-data tr:nth-child(even) td { background: #fcfcfc; }
"""

h = []
a = h.append

a('<!doctype html><html lang="pt-BR"><head><meta charset="utf-8"><title>Memorial Sanitário — ' + nome_projeto + '</title><style>' + CSS + '</style></head><body>')

# COVER
a('<div class="page page-cover"><div class="cover-container"><div class="cover-top-line"></div>')
a('<div class="cover-frame"><h1 class="cover-title">Memorial</h1><div class="cover-subtitle">Sanitário</div><div class="cover-subline"></div>')
a('<div class="cover-owner">' + proprietario + '</div></div>')
a('<div class="cover-bottom"><span style="color:#aaaaaa;font-size:9pt;">Esgoto Sanitário, Ventilação e Tratamento no Lote</span><div style="font-size:16pt;color:#f1c40f;font-weight:bold;">⌂</div></div>')
a('</div></div>')

# PAGE 2: INFORMAÇÕES
a('<div class="page"><h1 class="sec-title">1 INFORMAÇÕES DO PROJETO</h1>')
a('<table class="tb-info"><tr><th colspan="2">INFORMAÇÕES GERAIS</th></tr>')
a('<tr><td>Empreendimento</td><td>' + nome_projeto + '</td></tr>')
a('<tr><td>Proprietário / Cliente</td><td>' + proprietario + '</td></tr>')
a('<tr><td>Localização / Cidade</td><td>' + localizacao + '</td></tr>')
a('<tr><td>Número de pavimentos</td><td>2 Pavimentos</td></tr>')
a('<tr><td>Tipo de Edificação</td><td>Residencial Unifamiliar</td></tr>')
a('</table>')
a('<p>Profissional Desenhista: <b>' + RESP["nome"] + '</b> (' + RESP["titulo"] + ')</p>')

a('<h1 class="sec-title">2 NORMAS UTILIZADAS</h1>')
a('<ul><li><b>NBR 8160:1999</b> – Esgoto sanitário.</li><li><b>NBR 7229:1997</b> – Tanques sépticos.</li><li><b>NBR 13969:1997</b> – Tratamento complementar e disposição final.</li></ul>')

a('<h1 class="sec-title">3 TRATAMENTO NO LOTE (FOSSA / FILTRO / SUMIDOURO)</h1>')
a('<table class="tb-data"><tr><th>Unidade</th><th>Volume / Área</th><th>Parâmetros Adotados</th></tr>')
a('<tr><td>Fossa Séptica (NBR 7229)</td><td><b>' + str(fossa.get("volume_util_adotado_l", 2000)) + ' L</b></td><td>V = 1000 + N(CT + K Lf)</td></tr>')
a('<tr><td>Filtro Anaeróbio (NBR 13969)</td><td><b>' + str(filtro.get("volume_util_l", 1000)) + ' L</b></td><td>Leito filtrante brita nº 4 (h = 1.20m)</td></tr>')
a('<tr><td>Sumidouro / Infiltração</td><td><b>' + str(sumidouro.get("area_obtida_m2", 12.0)) + ' m²</b></td><td>A = CD / I (Poço circular DN 2.00m)</td></tr>')
a('</table></div></body></html>')

if not os.path.exists(SAIDA):
    os.makedirs(SAIDA)

nome_base = "Memorial_Sanitario_" + proprietario.replace(" ", "_").replace("&", "e")

# 1. HTML
html_path = os.path.join(SAIDA, nome_base + ".html")
f_html = codecs.open(html_path, "w", encoding="utf-8")
f_html.write("\n".join(h))
f_html.close()
print("HTML gerado:", html_path)

# 2. PDF
if pisa:
    try:
        pdf_path = os.path.join(SAIDA, nome_base + ".pdf")
        f_in = codecs.open(html_path, "r", encoding="utf-8")
        html_text = f_in.read()
        f_in.close()
        f_out = open(pdf_path, "wb")
        pisa.CreatePDF(html_text, dest=f_out)
        f_out.close()
        print("PDF gerado:", pdf_path)
    except Exception as ex_pdf:
        print("Aviso ao gerar PDF:", ex_pdf)

# 3. DOCX
if Document:
    docx_path = os.path.join(SAIDA, nome_base + ".docx")
    doc = Document()
    doc.add_heading("Memorial Sanitário — " + nome_projeto, level=0)
    doc.add_paragraph("Proprietário: " + proprietario)
    doc.add_paragraph("Localização: " + localizacao)
    doc.save(docx_path)
    print("DOCX gerado:", docx_path)
