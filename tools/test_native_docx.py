# -*- coding: utf-8 -*-
"""Generator for native Microsoft Word (.docx) calculation reports."""
import codecs
import json
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

if "RAIZ" in globals():
    RAIZ = globals()["RAIZ"]
elif "__file__" in globals():
    _this_dir = os.path.dirname(os.path.abspath(__file__))
    RAIZ = os.path.dirname(_this_dir) if os.path.basename(_this_dir) == "tools" else _this_dir
else:
    RAIZ = os.environ.get("HYDRO_PROJECT_ROOT", os.getcwd())
D = os.path.join(RAIZ, "data")
SAIDA = os.path.join(RAIZ, "memoriais")


def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml('<w:shd {0} w:fill="{1}"/>'.format(nsdecls('w'), fill_hex))
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml('<w:tcMar {0}><w:top w:w="{1}" w:type="dxa"/><w:bottom w:w="{2}" w:type="dxa"/><w:left w:w="{3}" w:type="dxa"/><w:right w:w="{4}" w:type="dxa"/></w:tcMar>'.format(
        nsdecls('w'), top, bottom, left, right))
    tcPr.append(tcMar)


def build_docx_report(titulo_disciplina, icone_emoji, dados_projeto, dados_tabela_pressao=None):
    doc = Document()
    
    # Page setup A4
    for sec in doc.sections:
        sec.page_width = Inches(8.27)
        sec.page_height = Inches(11.69)
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)
        
    proprietario = dados_projeto.get("proprietario", "Cliente do Projeto")
    nome_proj = dados_projeto.get("nome", "Residência Unifamiliar")
    cidade = dados_projeto.get("cidade", "Praia do Ervino - SFS")
    profissional = dados_projeto.get("profissional", "Thayná Barreiro")
    
    # COVER PAGE
    p_cov_space = doc.add_paragraph()
    p_cov_space.paragraph_format.space_before = Pt(80)
    
    p_title = doc.add_paragraph()
    r_t = p_title.add_run("Memorial\n")
    r_t.font.name = "Arial"
    r_t.font.size = Pt(44)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    r_sub = p_title.add_run(titulo_disciplina + " " + icone_emoji)
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(22)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    p_div = doc.add_paragraph()
    r_div = p_div.add_run("________________________________________")
    r_div.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p_div.paragraph_format.space_after = Pt(120)
    
    p_owner = doc.add_paragraph()
    r_o = p_owner.add_run(proprietario)
    r_o.font.name = "Segoe UI"
    r_o.font.size = Pt(18)
    r_o.font.bold = True
    r_o.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    
    doc.add_page_break()
    
    # SECTION 1: INFORMAÇÕES
    h1 = doc.add_heading("1 INFORMAÇÕES DO PROJETO", level=1)
    h1.style.font.color.rgb = RGBColor(0x0f, 0x3d, 0x5c)
    
    table_info = doc.add_table(rows=6, cols=2)
    table_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_info.style = 'Table Grid'
    
    hdr_cells = table_info.rows[0].cells
    hdr_cells[0].merge(hdr_cells[1])
    hdr_cells[0].text = "INFORMAÇÕES GERAIS DA EDIFICAÇÃO"
    set_cell_background(hdr_cells[0], "153D5C")
    hdr_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    
    info_data = [
        ("Empreendimento", nome_proj),
        ("Proprietário / Cliente", proprietario),
        ("Localização / Cidade", cidade),
        ("Número de Pavimentos", "2 Pavimentos"),
        ("Profissional Desenhista", profissional)
    ]
    
    for idx, (k, v) in enumerate(info_data, start=1):
        row_cells = table_info.rows[idx].cells
        row_cells[0].text = k
        row_cells[1].text = v
        set_cell_background(row_cells[0], "F2F4F7")
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_margins(row_cells[0], 80, 80, 120, 120)
        set_cell_margins(row_cells[1], 80, 80, 120, 120)
        
    p_space = doc.add_paragraph()
    
    # SECTION 2: NORMAS
    h2 = doc.add_heading("2 NORMAS TÉCNICAS APLICÁVEIS", level=1)
    h2.style.font.color.rgb = RGBColor(0x0f, 0x3d, 0x5c)
    
    normas = [
        "NBR 5626:2020 – Sistemas Prediais de Água Fria e Água Quente (Projeto, Execução e Manutenção)",
        "NBR 8160:1999 – Sistemas Prediais de Esgoto Sanitário e Ventilação",
        "NBR 7229:1997 / NBR 13969:1997 – Tanques Sépticos e Tratamento no Lote",
        "NBR 10844:1989 / DTU 60.11 – Instalações Prediais de Águas Pluviais",
        "NBR 5674:2012 – Manutenção de Edificações"
    ]
    for n in normas:
        doc.add_paragraph(n, style='List Bullet')
        
    # SECTION 3: MEMORIAL DE CÁLCULO / VERIFICAÇÃO DE PRESSÃO
    if dados_tabela_pressao:
        doc.add_page_break()
        h3 = doc.add_heading("3 MEMORIAL DE CÁLCULO E VERIFICAÇÃO DE PRESSÃO", level=1)
        h3.style.font.color.rgb = RGBColor(0x0f, 0x3d, 0x5c)
        
        table_p = doc.add_table(rows=len(dados_tabela_pressao) + 1, cols=9)
        table_p.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_p.style = 'Table Grid'
        
        headers_p = ["Trecho", "Σ P", "Q (L/s)", "DN (mm)", "V (m/s)", "H (m)", "L (m)", "Pfinal (mca)", "Preq (mca)"]
        for col_idx, text in enumerate(headers_p):
            cell = table_p.rows[0].cells[col_idx]
            cell.text = text
            set_cell_background(cell, "0F3D5C")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)
            
        for row_idx, data in enumerate(dados_tabela_pressao, start=1):
            row_cells = table_p.rows[row_idx].cells
            vals = [
                data.get("nome", "Trecho"),
                str(data.get("peso", 0.4)),
                str(data.get("Q_ls", 0.2)),
                str(data.get("dn_mm", 25)),
                str(data.get("v_ms", 1.2)),
                str(data.get("h_fin", 3.5)),
                str(data.get("l_real", 2.5)),
                str(round(data.get("p_fin", 3.33), 2)),
                str(data.get("p_req", 1.0))
            ]
            for col_idx, val in enumerate(vals):
                row_cells[col_idx].text = val
                row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(8.5)
                set_cell_margins(row_cells[col_idx], 50, 50, 60, 60)
                
    if not os.path.exists(SAIDA):
        os.makedirs(SAIDA)
        
    file_path = os.path.join(SAIDA, "Memorial_" + titulo_disciplina.replace(" ", "_") + "_" + proprietario.replace(" ", "_").replace("&", "e") + ".docx")
    doc.save(file_path)
    print("DOCX generated successfully:", file_path)
    return file_path

# Test build
if __name__ == "__main__":
    build_docx_report("Hidráulico", "💧", {"proprietario": "Henrique e Suelen", "nome": "Casa A&R", "cidade": "Porto Alegre"}, [
        {"nome": "1-2 Barrilete", "peso": 5.5, "Q_ls": 0.70, "dn_mm": 32, "v_ms": 0.88, "h_fin": 3.50, "l_real": 3.5, "p_fin": 5.93, "p_req": 1.0}
    ])
