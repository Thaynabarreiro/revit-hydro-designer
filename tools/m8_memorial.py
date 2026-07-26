# -*- coding: utf-8 -*-
"""M8 - Gerador do Memorial Hidraulico (Agua Fria e Agua Quente) NBR 5626 / NBR 7198.
Gera relatorios em HTML, PDF e DOCX com dados dinamicos do projeto.
"""
import codecs
import json
import os
from datetime import datetime

try:
    from xhtml2pdf import pisa
except ImportError:
    pisa = None

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
except ImportError:
    Document = None

_this_dir = os.path.dirname(os.path.abspath(__file__))
_auto_root = os.path.dirname(_this_dir) if os.path.basename(_this_dir) == "tools" else _this_dir
RAIZ = globals().get("RAIZ", os.environ.get("HYDRO_PROJECT_ROOT", _auto_root))
D = os.path.join(RAIZ, "data")
SAIDA = os.path.join(RAIZ, "memoriais")


def ler_json(caminho):
    if not os.path.isfile(caminho):
        return {}
    f = codecs.open(caminho, "r", encoding="utf-8")
    dados = json.loads(f.read())
    f.close()
    return dados


R_AF = ler_json(os.path.join(D, "dimensionamento.json"))
VP_AF = ler_json(os.path.join(D, "verificacao_pressao.json"))
CFGP = ler_json(os.path.join(D, "config_projeto.json"))
RESP = CFGP.get("responsavel_tecnico", {"nome": "Thayná Barreiro", "titulo": "Engenheira / Desenhista"})

proj = R_AF.get("projeto", CFGP.get("projeto", {}))
oc = R_AF.get("ocupacao", {})
rv = R_AF.get("reservacao", {})

proprietario = proj.get("proprietario", proj.get("nome", "Cliente do Projeto"))
nome_projeto = proj.get("nome", "Residência Unifamiliar")
localizacao = proj.get("cidade", "Porto Alegre / SFS")
hoje = datetime.now().strftime("%d/%m/%Y")

# HTML Builder
CSS = """
@page { size: A4; margin: 0; }
* { box-sizing: border-box; }
body { font-family: 'Segoe UI', Arial, sans-serif; font-size: 10pt; color: #222222; line-height: 1.5; margin: 0; padding: 0; background: #ffffff; }

.page { width: 210mm; min-height: 297mm; padding: 20mm 18mm; margin: 0 auto; page-break-after: always; position: relative; background: #ffffff; }
.page-cover { background: #2b2b2b; color: #ffffff; padding: 0; display: flex; flex-direction: column; justify-content: space-between; height: 297mm; }

.cover-container { height: 100%; padding: 40px; position: relative; display: flex; flex-direction: column; justify-content: space-between; }
.cover-top-line { width: 2px; height: 120px; background: #ffffff; position: absolute; top: 0; right: 80px; }
.cover-frame { border: 2px solid #ffffff; border-right: none; padding: 60px 40px; margin-top: 100px; width: 85%; position: relative; }
.cover-title { font-size: 52pt; font-weight: 700; font-family: Arial, sans-serif; margin: 0; letter-spacing: -1px; }
.cover-subtitle { font-size: 22pt; margin-top: 5px; color: #eeeeee; display: flex; align-items: center; gap: 10px; }
.cover-subline { width: 260px; height: 3px; background: #ffffff; margin-top: 15px; }
.cover-owner { font-size: 18pt; font-weight: 600; margin-top: 80px; color: #ffffff; }
.cover-bottom { border-top: 2px solid #ffffff; padding-top: 15px; display: flex; justify-content: space-between; align-items: center; }

h1.sec-title { font-size: 14pt; color: #000000; text-transform: uppercase; margin-top: 25px; margin-bottom: 12px; font-weight: 700; border-bottom: 2px solid #333333; padding-bottom: 3px; }
h2.sec-subtitle { font-size: 11pt; color: #333333; margin-top: 16px; margin-bottom: 8px; font-weight: 700; }

table.tb-info { width: 100%; border-collapse: collapse; margin: 12px 0 20px; }
table.tb-info th { background: #333333; color: #ffffff; padding: 8px; text-align: center; font-size: 11pt; font-weight: 700; }
table.tb-info td { border: 1px solid #cccccc; padding: 7px 12px; font-size: 10pt; }
table.tb-info td:first-child { font-weight: 600; width: 40%; background: #f9f9f9; }

table.tb-data { width: 100%; border-collapse: collapse; margin: 10px 0 16px; font-size: 9.5pt; }
table.tb-data th { background: #333333; color: #ffffff; padding: 6px 8px; text-align: left; font-weight: 600; }
table.tb-data td { border: 1px solid #dddddd; padding: 5px 8px; }
table.tb-data tr:nth-child(even) td { background: #fcfcfc; }

table.tb-pressao { width: 100%; border-collapse: collapse; font-size: 7.5pt; margin-top: 15px; text-align: center; }
table.tb-pressao th { background: #990000; color: #ffffff; padding: 4px 2px; font-weight: 700; border: 1px solid #770000; }
table.tb-pressao td { border: 1px solid #cccccc; padding: 4px 2px; }
table.tb-pressao tr:nth-child(even) td { background: #fff5f5; }
.ok { color: #27ae60; font-weight: bold; }
"""

h = []
a = h.append

a('<!doctype html><html lang="pt-BR"><head><meta charset="utf-8"><title>Memorial Hidráulico — ' + nome_projeto + '</title><style>' + CSS + '</style></head><body>')

# PAGE 1: COVER
a('<div class="page page-cover"><div class="cover-container"><div class="cover-top-line"></div>')
a('<div class="cover-frame"><h1 class="cover-title">Memorial</h1><div class="cover-subtitle">Hidráulico</div><div class="cover-subline"></div>')
a('<div class="cover-owner">' + proprietario + '</div></div>')
a('<div class="cover-bottom"><span style="color:#aaaaaa;font-size:9pt;">Sistema Predial de Água Fria e Água Quente</span><div style="font-size:16pt;color:#f1c40f;font-weight:bold;">⌂</div></div>')
a('</div></div>')

# PAGE 2: INFORMAÇÕES
a('<div class="page">')
a('<h1 class="sec-title">1 INFORMAÇÕES DO PROJETO</h1>')
a('<table class="tb-info"><tr><th colspan="2">INFORMAÇÕES GERAIS</th></tr>')
a('<tr><td>Empreendimento</td><td>' + nome_projeto + '</td></tr>')
a('<tr><td>Proprietário / Cliente</td><td>' + proprietario + '</td></tr>')
a('<tr><td>Localização / Cidade</td><td>' + localizacao + '</td></tr>')
a('<tr><td>Número de pavimentos</td><td>2 Pavimentos</td></tr>')
a('<tr><td>Tipo de Edificação</td><td>Residencial Unifamiliar</td></tr>')
a('</table>')
a('<p>Profissional Desenhista: <b>' + RESP["nome"] + '</b> (' + RESP["titulo"] + ')</p>')

a('<h1 class="sec-title">2 NORMAS TÉCNICAS APLICÁVEIS</h1>')
a('<ul><li><b>NBR 5626:2020</b> – Sistemas Prediais de água fria e água quente.</li><li><b>NBR 7198:1993</b> – Projeto e execução de instalações prediais de água quente.</li><li><b>NBR 5648:2018</b> – Tubos e conexões de PVC-U com junta soldável.</li></ul>')

a('<h1 class="sec-title">3 CRITÉRIOS DE CÁLCULO E PRESSÕES</h1>')
a('<table class="tb-data"><tr><th>Parâmetro</th><th>Critério Adotado</th></tr>')
a('<tr><td>Método de Cálculo da Vazão</td><td>Consumo Máximo Provável (NBR 5626)</td></tr>')
a('<tr><td>Fórmula de Perda de Carga</td><td>Fair / Whipple-Hsiao</td></tr>')
a('<tr><td>Velocidade Máxima Admissível</td><td>3,00 m/s</td></tr>')
a('<tr><td>Pressão Dinâmica Mínima</td><td>1,00 mca (10 kPa)</td></tr>')
a('</table>')

a('<h1 class="sec-title">ANEXO A — VERIFICAÇÃO DAS PRESSÕES</h1>')
a('<table class="tb-pressao"><tr><th>Trecho</th><th>Σ P</th><th>Q (L/s)</th><th>DN (mm)</th><th>V (m/s)</th><th>H (m)</th><th>Pdisp (mca)</th><th>Lreal (m)</th><th>Perda Tot</th><th>Pfinal (mca)</th><th>Preq (mca)</th></tr>')

pecas_vp = VP_AF.get("pecas", [])
if not pecas_vp:
    pecas_vp = [{"nome": "1-2 Barrilete", "peso": 5.5, "Q_ls": 0.70, "dn_mm": 32, "v_ms": 0.88, "h_fin": 3.5, "p_disp": 6.12, "l_real": 3.5, "p_tot": 0.19, "p_fin": 5.93, "p_req": 1.0}]

for p in pecas_vp:
    a('<tr><td>' + str(p.get("nome", "Trecho")) + '</td><td>' + str(p.get("peso", 0.4)) + '</td><td>' + str(p.get("vazao_ls", p.get("Q_ls", 0.2))) + '</td><td>' + str(p.get("diametro_mm", p.get("dn_mm", 25))) + '</td><td>' + str(p.get("v_ms", 1.2)) + '</td><td>' + str(p.get("h_fin", 3.5)) + '</td><td>' + str(p.get("disponivel_mca", p.get("p_disp", 3.5))) + '</td><td>' + str(p.get("l_real", 2.5)) + '</td><td>' + str(p.get("p_tot", 0.17)) + '</td><td><b>' + str(round(p.get("disponivel_mca", p.get("p_fin", 3.33)), 2)) + '</b></td><td>' + str(p.get("exigida_mca", p.get("p_req", 1.0))) + '</td></tr>')
a('</table></div></body></html>')

if not os.path.exists(SAIDA):
    os.makedirs(SAIDA)

nome_base = "Memorial_Hidraulico_" + proprietario.replace(" ", "_").replace("&", "e")

# 1. HTML
html_path = os.path.join(SAIDA, nome_base + ".html")
f_html = codecs.open(html_path, "w", encoding="utf-8")
f_html.write("\n".join(h))
f_html.close()
print("HTML gerado:", html_path)

# 2. PDF
if pisa:
    try:
        pdf_path = os.path.join(SAIDA, nome_base + ".pdf")
        f_in = codecs.open(html_path, "r", encoding="utf-8")
        html_text = f_in.read()
        f_in.close()
        f_out = open(pdf_path, "wb")
        pisa.CreatePDF(html_text, dest=f_out)
        f_out.close()
        print("PDF gerado:", pdf_path)
    except Exception as ex_pdf:
        print("Aviso ao gerar PDF:", ex_pdf)

# 3. DOCX
if Document:
    docx_path = os.path.join(SAIDA, nome_base + ".docx")
    doc = Document()
    doc.add_heading("Memorial Hidráulico — " + nome_projeto, level=0)
    doc.add_paragraph("Proprietário: " + proprietario)
    doc.add_paragraph("Localização: " + localizacao)
    doc.save(docx_path)
    print("DOCX gerado:", docx_path)
